"""Guide content API — CRUD for user guide sections, blocks, screenshots, and export."""

import os
import re
import uuid

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import select, func as sa_func
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db
from app.models.guide_content import GuideBlock, GuideSection
from app.models.user import User
from app.services.auth_service import get_current_user

router = APIRouter(prefix="/api/guide", tags=["guide"])

UPLOAD_DIR = "/app/uploads/guide"


# ── Schemas ───────────────────────────────────────────────────
class SectionCreate(BaseModel):
    title: str

class SectionUpdate(BaseModel):
    title: str | None = None
    order_index: int | None = None
    is_visible: bool | None = None

class BlockCreate(BaseModel):
    block_type: str
    content_json: dict | None = None

class BlockUpdate(BaseModel):
    content_json: dict | None = None
    order_index: int | None = None
    is_visible: bool | None = None

class ReorderRequest(BaseModel):
    ids: list[str]


# ── Helpers ───────────────────────────────────────────────────
def _slugify(text: str) -> str:
    s = text.lower().strip()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_]+", "-", s)
    return re.sub(r"-+", "-", s).strip("-")[:200]


def _section_dict(s: GuideSection, blocks: list[dict] | None = None) -> dict:
    d = {
        "id": str(s.id), "order_index": s.order_index, "title": s.title,
        "slug": s.slug, "is_visible": s.is_visible,
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "updated_at": s.updated_at.isoformat() if s.updated_at else None,
    }
    if blocks is not None:
        d["blocks"] = blocks
    return d


def _block_dict(b: GuideBlock) -> dict:
    return {
        "id": str(b.id), "section_id": str(b.section_id), "order_index": b.order_index,
        "block_type": b.block_type, "content_json": b.content_json or {},
        "is_visible": b.is_visible,
        "created_at": b.created_at.isoformat() if b.created_at else None,
    }


# ── PUBLIC: get full guide ────────────────────────────────────
@router.get("")
async def get_guide(db: AsyncSession = Depends(get_db)):
    """Get the full guide content — all visible sections with visible blocks, ordered."""
    sections = (await db.execute(
        select(GuideSection).where(GuideSection.is_visible == True)  # noqa: E712
        .order_by(GuideSection.order_index)
    )).scalars().all()

    result = []
    for s in sections:
        blocks = (await db.execute(
            select(GuideBlock).where(
                GuideBlock.section_id == s.id,
                GuideBlock.is_visible == True,  # noqa: E712
            ).order_by(GuideBlock.order_index)
        )).scalars().all()
        result.append(_section_dict(s, [_block_dict(b) for b in blocks]))

    return {"sections": result}


# ── ADMIN: get all including hidden ───────────────────────────
@router.get("/admin")
async def get_guide_admin(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sections = (await db.execute(
        select(GuideSection).order_by(GuideSection.order_index)
    )).scalars().all()

    result = []
    for s in sections:
        blocks = (await db.execute(
            select(GuideBlock).where(GuideBlock.section_id == s.id)
            .order_by(GuideBlock.order_index)
        )).scalars().all()
        result.append(_section_dict(s, [_block_dict(b) for b in blocks]))

    return {"sections": result}


# ── SECTIONS CRUD ─────────────────────────────────────────────
@router.post("/sections")
async def create_section(
    body: SectionCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    max_order = (await db.execute(
        select(sa_func.coalesce(sa_func.max(GuideSection.order_index), -1))
    )).scalar() or -1

    slug = _slugify(body.title)
    # Ensure unique slug
    existing = (await db.execute(select(GuideSection).where(GuideSection.slug == slug))).scalar_one_or_none()
    if existing:
        slug = f"{slug}-{uuid.uuid4().hex[:4]}"

    section = GuideSection(title=body.title, slug=slug, order_index=max_order + 1)
    db.add(section)
    await db.flush()
    return _section_dict(section, [])


@router.put("/sections/{section_id}")
async def update_section(
    section_id: uuid.UUID, body: SectionUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    section = (await db.execute(select(GuideSection).where(GuideSection.id == section_id))).scalar_one_or_none()
    if not section:
        raise HTTPException(404, "Section not found")
    updates = body.model_dump(exclude_unset=True)
    if "title" in updates:
        section.title = updates["title"]
        section.slug = _slugify(updates["title"])
    if "is_visible" in updates:
        section.is_visible = updates["is_visible"]
    if "order_index" in updates:
        section.order_index = updates["order_index"]
    await db.flush()
    # Build response from known fields to avoid lazy loads
    return {
        "id": str(section.id), "order_index": section.order_index,
        "title": section.title, "slug": section.slug, "is_visible": section.is_visible,
    }


@router.delete("/sections/{section_id}")
async def delete_section(
    section_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    section = (await db.execute(select(GuideSection).where(GuideSection.id == section_id))).scalar_one_or_none()
    if not section:
        raise HTTPException(404, "Section not found")
    # Delete screenshot files from blocks
    blocks = (await db.execute(select(GuideBlock).where(GuideBlock.section_id == section_id))).scalars().all()
    for b in blocks:
        if b.block_type == "screenshot" and b.content_json and b.content_json.get("image_path"):
            _try_delete_file(b.content_json["image_path"])
        await db.delete(b)
    await db.delete(section)
    return {"deleted": True}


@router.post("/sections/reorder")
async def reorder_sections(
    body: ReorderRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    for idx, sid in enumerate(body.ids):
        section = (await db.execute(select(GuideSection).where(GuideSection.id == uuid.UUID(sid)))).scalar_one_or_none()
        if section:
            section.order_index = idx
    await db.flush()
    return {"reordered": len(body.ids)}


# ── BLOCKS CRUD ───────────────────────────────────────────────
@router.post("/sections/{section_id}/blocks")
async def create_block(
    section_id: uuid.UUID, body: BlockCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    section = (await db.execute(select(GuideSection).where(GuideSection.id == section_id))).scalar_one_or_none()
    if not section:
        raise HTTPException(404, "Section not found")

    max_order = (await db.execute(
        select(sa_func.coalesce(sa_func.max(GuideBlock.order_index), -1))
        .where(GuideBlock.section_id == section_id)
    )).scalar() or -1

    block = GuideBlock(
        section_id=section_id, block_type=body.block_type,
        content_json=body.content_json or _default_content(body.block_type),
        order_index=max_order + 1,
    )
    db.add(block)
    await db.flush()
    return _block_dict(block)


@router.put("/sections/{section_id}/blocks/{block_id}")
async def update_block(
    section_id: uuid.UUID, block_id: uuid.UUID, body: BlockUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    block = (await db.execute(
        select(GuideBlock).where(GuideBlock.id == block_id, GuideBlock.section_id == section_id)
    )).scalar_one_or_none()
    if not block:
        raise HTTPException(404, "Block not found")
    updates = body.model_dump(exclude_unset=True)
    if "content_json" in updates:
        block.content_json = updates["content_json"]
    if "order_index" in updates:
        block.order_index = updates["order_index"]
    if "is_visible" in updates:
        block.is_visible = updates["is_visible"]
    await db.flush()
    return {
        "id": str(block.id), "section_id": str(block.section_id),
        "order_index": block.order_index, "block_type": block.block_type,
        "content_json": block.content_json or {}, "is_visible": block.is_visible,
    }


@router.delete("/sections/{section_id}/blocks/{block_id}")
async def delete_block(
    section_id: uuid.UUID, block_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    block = (await db.execute(
        select(GuideBlock).where(GuideBlock.id == block_id, GuideBlock.section_id == section_id)
    )).scalar_one_or_none()
    if not block:
        raise HTTPException(404, "Block not found")
    if block.block_type == "screenshot" and block.content_json and block.content_json.get("image_path"):
        _try_delete_file(block.content_json["image_path"])
    await db.delete(block)
    return {"deleted": True}


@router.post("/sections/{section_id}/blocks/reorder")
async def reorder_blocks(
    section_id: uuid.UUID, body: ReorderRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    for idx, bid in enumerate(body.ids):
        block = (await db.execute(
            select(GuideBlock).where(GuideBlock.id == uuid.UUID(bid), GuideBlock.section_id == section_id)
        )).scalar_one_or_none()
        if block:
            block.order_index = idx
    await db.flush()
    return {"reordered": len(body.ids)}


# ── SCREENSHOTS ───────────────────────────────────────────────
@router.post("/screenshots/upload")
async def upload_screenshot(
    file: UploadFile,
    current_user: User = Depends(get_current_user),
):
    if file.content_type not in ("image/png", "image/jpeg", "image/jpg"):
        raise HTTPException(400, "Only PNG and JPG")
    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(400, "Max 5MB")

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    safe = re.sub(r"[^\w.\-]", "_", file.filename or "image.png")
    filename = f"{uuid.uuid4().hex[:8]}_{safe}"
    save_path = os.path.join(UPLOAD_DIR, filename)
    with open(save_path, "wb") as f:
        f.write(content)

    w, h = 0, 0
    try:
        from PIL import Image
        img = Image.open(save_path)
        w, h = img.size
    except Exception:
        pass

    return {"image_path": f"/uploads/guide/{filename}", "filename": filename, "width": w, "height": h}


# ── SEED ──────────────────────────────────────────────────────
@router.post("/seed")
async def seed_guide(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Seed default guide content. Only works if no sections exist."""
    existing = (await db.execute(select(sa_func.count(GuideSection.id)))).scalar()
    if existing and existing > 0:
        return {"seeded": False, "message": "Guide already has content"}

    from app.services.guide_seed import seed_default_guide
    count = await seed_default_guide(db)
    return {"seeded": True, "sections": count}


# ── EXPORT ────────────────────────────────────────────────────
@router.post("/export/pdf")
async def export_pdf(db: AsyncSession = Depends(get_db)):
    sections = (await db.execute(
        select(GuideSection).where(GuideSection.is_visible == True).order_by(GuideSection.order_index)  # noqa
    )).scalars().all()

    html_parts = ['<html><head><meta charset="utf-8"><style>',
        'body{font-family:Arial,sans-serif;max-width:800px;margin:40px auto;color:#1f2937;font-size:11pt;line-height:1.6}',
        'h1{color:#00338D;font-size:22pt;border-bottom:3px solid #00338D;padding-bottom:8pt;margin-top:30pt}',
        'h2{color:#00338D;font-size:16pt;margin-top:24pt}h3{font-size:13pt;margin-top:18pt}',
        'p{margin:6pt 0}img{max-width:100%;border-radius:6pt;border:1pt solid #e5e7eb;margin:12pt 0}',
        '.tip{background:#EFF6FF;border:1pt solid #BFDBFE;border-radius:6pt;padding:10pt 14pt;margin:10pt 0}',
        '.warn{background:#FFFBEB;border:1pt solid #FDE68A;border-radius:6pt;padding:10pt 14pt;margin:10pt 0}',
        'table{width:100%;border-collapse:collapse;margin:10pt 0}th{background:#F3F4F6;text-align:left;padding:6pt 10pt;border-bottom:2pt solid #D1D5DB}',
        'td{padding:6pt 10pt;border-bottom:1pt solid #E5E7EB}code{background:#F3F4F6;padding:1pt 4pt;border-radius:3pt;font-size:9pt}',
        '.step{display:flex;gap:8pt;margin:6pt 0;align-items:flex-start}',
        '.step-num{background:#3B82F6;color:white;border-radius:50%;width:20pt;height:20pt;display:flex;align-items:center;justify-content:center;font-size:9pt;font-weight:bold;flex-shrink:0}',
        '</style></head><body>',
        '<h1>Slides Generator by KPMG &mdash; User Guide</h1>']

    for s in sections:
        html_parts.append(f'<h2>{s.title}</h2>')
        blocks = (await db.execute(
            select(GuideBlock).where(GuideBlock.section_id == s.id, GuideBlock.is_visible == True)  # noqa
            .order_by(GuideBlock.order_index)
        )).scalars().all()
        for b in blocks:
            c = b.content_json or {}
            if b.block_type == "heading":
                lvl = c.get("level", 2)
                html_parts.append(f'<h{lvl}>{c.get("text", "")}</h{lvl}>')
            elif b.block_type == "paragraph":
                html_parts.append(f'<p>{c.get("text", "")}</p>')
            elif b.block_type == "screenshot" and c.get("image_path"):
                img_path = "/app" + c["image_path"] if c["image_path"].startswith("/uploads") else c["image_path"]
                html_parts.append(f'<img src="file://{img_path}" alt="{c.get("alt", "")}" />')
                if c.get("caption"):
                    html_parts.append(f'<p style="font-size:9pt;color:#9CA3AF;text-align:center;font-style:italic">{c["caption"]}</p>')
            elif b.block_type == "tip":
                html_parts.append(f'<div class="tip"><strong>Tip:</strong> {c.get("text", "")}</div>')
            elif b.block_type == "warning":
                html_parts.append(f'<div class="warn"><strong>Warning:</strong> {c.get("text", "")}</div>')
            elif b.block_type == "steps":
                for i, item in enumerate(c.get("items", []), 1):
                    html_parts.append(f'<div class="step"><div class="step-num">{i}</div><span>{item}</span></div>')
            elif b.block_type == "shortcut_table":
                html_parts.append('<table><tr><th>Shortcut</th><th>Action</th></tr>')
                for row in c.get("rows", []):
                    html_parts.append(f'<tr><td><code>{row.get("key", "")}</code></td><td>{row.get("action", "")}</td></tr>')
                html_parts.append('</table>')
            elif b.block_type == "divider":
                html_parts.append('<hr style="border:none;border-top:1pt solid #E5E7EB;margin:16pt 0">')

    html_parts.append('</body></html>')
    html_str = "\n".join(html_parts)

    import tempfile
    output_dir = tempfile.mkdtemp()
    filepath = os.path.join(output_dir, "KPMG_Slide_Generator_Guide.pdf")
    from weasyprint import HTML
    HTML(string=html_str).write_pdf(filepath)
    return FileResponse(filepath, media_type="application/pdf", filename="KPMG_Slide_Generator_Guide.pdf")


@router.post("/export/docx")
async def export_docx(db: AsyncSession = Depends(get_db)):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    import tempfile

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Slides Generator by KPMG — User Guide", level=0)
    doc.add_paragraph("")

    sections = (await db.execute(
        select(GuideSection).where(GuideSection.is_visible == True).order_by(GuideSection.order_index)  # noqa
    )).scalars().all()

    for s in sections:
        doc.add_heading(s.title, level=1)
        blocks = (await db.execute(
            select(GuideBlock).where(GuideBlock.section_id == s.id, GuideBlock.is_visible == True)  # noqa
            .order_by(GuideBlock.order_index)
        )).scalars().all()
        for b in blocks:
            c = b.content_json or {}
            if b.block_type == "heading":
                doc.add_heading(c.get("text", ""), level=c.get("level", 2))
            elif b.block_type == "paragraph":
                doc.add_paragraph(c.get("text", ""))
            elif b.block_type == "screenshot" and c.get("image_path"):
                img_path = "/app" + c["image_path"] if c["image_path"].startswith("/uploads") else c["image_path"]
                try:
                    doc.add_picture(img_path, width=Inches(6))
                except Exception:
                    pass
                if c.get("caption"):
                    p = doc.add_paragraph(c["caption"])
                    p.runs[0].italic = True
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            elif b.block_type == "tip":
                p = doc.add_paragraph()
                run = p.add_run(f"Tip: {c.get('text', '')}")
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            elif b.block_type == "warning":
                p = doc.add_paragraph()
                run = p.add_run(f"Warning: {c.get('text', '')}")
                run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            elif b.block_type == "steps":
                for i, item in enumerate(c.get("items", []), 1):
                    doc.add_paragraph(f"{i}. {item}")
            elif b.block_type == "shortcut_table":
                rows = c.get("rows", [])
                if rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=2)
                    table.style = "Light Grid Accent 1"
                    table.cell(0, 0).text = "Shortcut"
                    table.cell(0, 1).text = "Action"
                    for i, row in enumerate(rows, 1):
                        table.cell(i, 0).text = row.get("key", "")
                        table.cell(i, 1).text = row.get("action", "")

    output_dir = tempfile.mkdtemp()
    filepath = os.path.join(output_dir, "KPMG_Slide_Generator_Guide.docx")
    doc.save(filepath)
    return FileResponse(filepath, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename="KPMG_Slide_Generator_Guide.docx")


# ── Utils ─────────────────────────────────────────────────────
def _try_delete_file(image_path: str):
    if not image_path:
        return
    full = os.path.join("/app", image_path.lstrip("/"))
    if os.path.exists(full):
        try:
            os.remove(full)
        except OSError:
            pass


def _default_content(block_type: str) -> dict:
    defaults = {
        "heading": {"text": "New Heading", "level": 2},
        "paragraph": {"text": ""},
        "screenshot": {"image_path": None, "caption": "", "alt": "", "annotations": []},
        "tip": {"text": ""},
        "warning": {"text": ""},
        "steps": {"items": ["Step 1"]},
        "shortcut_table": {"rows": [{"key": "Ctrl+Z", "action": "Undo"}]},
        "divider": {},
    }
    return defaults.get(block_type, {})
